import os
import json
import io
import base64
import email
from flask import Flask, request, jsonify, send_from_directory
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

CRITERIOS = {
    "Cobana": [
        {"texto": "C Appearance 4 Qlty/Cnd 2 Storage", "calificacion": "REGULAR"},
        {"texto": "C Appearance 2 Qlty/Cnd 1.3 Storage", "calificacion": "ACEPTABLE"},
        {"texto": "B Appearance 3 Qlty/Cnd 1.8 Storage", "calificacion": "BUENO"},
        {"texto": "C Appearance 4 Qlty/Cnd 2.8 Storage", "calificacion": "POBRE"},
        {"texto": "C Appearance 3 Qlty/Cnd 2.3 Storage", "calificacion": "POBRE"},
    ],
    "Oppy": [
        {"texto": "Decay % < 1 - Serious % < 5 - Total % < 10", "calificacion": "BUENO"},
        {"texto": "Decay % > 1 - Serious % < 5 - Total % < 10", "calificacion": "POBRE"},
        {"texto": "Decay % < 1 - Serious % > 5 - Total % < 10", "calificacion": "POBRE"},
        {"texto": "Decay % < 1 - Serious % < 5 - Total % > 10", "calificacion": "POBRE"},
    ],
}

def extraer_texto_pdf(file_bytes):
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text.strip()

def extraer_texto_word(file_bytes):
    import docx
    doc = docx.Document(io.BytesIO(file_bytes))
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + " | ".join([cell.text for cell in row.cells])
    return text.strip()

def extraer_texto_excel(file_bytes):
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True)
    text = ""
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        text += f"\nHoja: {sheet}\n"
        for row in ws.iter_rows(values_only=True):
            fila = " | ".join([str(c) if c is not None else "" for c in row])
            if fila.strip():
                text += fila + "\n"
    return text.strip()

def extraer_texto_email(file_bytes, filename):
    import re
    ext = filename.lower().split('.')[-1]
    
    if ext == 'msg':
        import extract_msg
        msg = extract_msg.Message(io.BytesIO(file_bytes))
        text = f"Asunto: {msg.subject or ''}\n"
        text += f"De: {msg.sender or ''}\n"
        text += f"Fecha: {msg.date or ''}\n\n"
        text += msg.body or ""
        return text.strip()
    else:
        msg = email.message_from_bytes(file_bytes)
        text = f"Asunto: {msg.get('Subject', '')}\n"
        text += f"De: {msg.get('From', '')}\n"
        text += f"Fecha: {msg.get('Date', '')}\n\n"
        for part in msg.walk():
            ct = part.get_content_type()
            if ct == "text/plain":
                payload = part.get_payload(decode=True)
                if payload:
                    text += payload.decode("utf-8", errors="ignore")
        return text.strip() 

def extraer_contenido(file_bytes, filename):
    ext = filename.lower().split('.')[-1]
    if ext == 'pdf':
        return {"tipo": "texto", "contenido": extraer_texto_pdf(file_bytes)}
    elif ext in ['doc', 'docx']:
        return {"tipo": "texto", "contenido": extraer_texto_word(file_bytes)}
    elif ext in ['xls', 'xlsx']:
        return {"tipo": "texto", "contenido": extraer_texto_excel(file_bytes)}
    elif ext in ['eml', 'msg']:
    return {"tipo": "texto", "contenido": extraer_texto_email(file_bytes, filename)}
    elif ext in ['txt', 'csv']:
        return {"tipo": "texto", "contenido": file_bytes.decode("utf-8", errors="ignore")}
    elif ext in ['png', 'jpg', 'jpeg', 'gif', 'webp', 'bmp']:
        b64 = base64.b64encode(file_bytes).decode("utf-8")
        mime = f"image/{ext}" if ext != 'jpg' else "image/jpeg"
        return {"tipo": "imagen", "contenido": b64, "mime": mime}
    else:
        return {"tipo": "error", "contenido": "Formato no soportado"}

@app.route("/")
def index():
    return send_from_directory(".", "index.html")

@app.route("/analizar", methods=["POST"])
def analizar():
    try:
        cliente = request.form.get("cliente")
        archivo = request.files.get("archivo")
        if not cliente or not archivo:
            return jsonify({"error": "Faltan datos"}), 400

        file_bytes = archivo.read()
        filename = archivo.filename
        resultado_extraccion = extraer_contenido(file_bytes, filename)

        if resultado_extraccion["tipo"] == "error":
            return jsonify({"error": resultado_extraccion["contenido"]}), 400

        criterios = CRITERIOS.get(cliente, [])
        criterios_texto = "\n".join(
            [f'{i+1}. "{c["texto"]}" → {c["calificacion"]}' for i, c in enumerate(criterios)]
        )

        prompt_base = f"""Sos un agente especializado en análisis de informes de calidad (QC) de exportaciones de cítricos de URUD'OR S.A.

Tu tarea es analizar el informe del cliente {cliente} y determinar la calificación según estos criterios:

CRITERIOS PARA {cliente.upper()}:
{criterios_texto}

Respondé ÚNICAMENTE con un objeto JSON válido con esta estructura:
{{
  "contenedor": "código del contenedor",
  "buque": "nombre del buque",
  "variedad": "variedad de fruta",
  "arrival": "fecha de llegada",
  "calificacion": "BUENO|ACEPTABLE|REGULAR|POBRE",
  "scores_detectados": "valores clave encontrados en el informe que determinaron la calificacion",
  "razonamiento": "explicación breve en español de por qué corresponde esa calificación según los criterios"
}}

No incluyas texto fuera del JSON. No uses markdown ni backticks."""

        if resultado_extraccion["tipo"] == "texto":
            contenido = resultado_extraccion["contenido"]
            if not contenido or len(contenido.strip()) < 30:
                return jsonify({"error": "No se pudo extraer contenido del archivo"}), 400
            messages = [{"role": "user", "content": prompt_base + f"\n\nCONTENIDO DEL INFORME:\n{contenido}"}]

        elif resultado_extraccion["tipo"] == "imagen":
            messages = [{"role": "user", "content": [
                {"type": "text", "text": prompt_base + "\n\nEl informe se adjunta como imagen:"},
                {"type": "image_url", "image_url": {"url": f"data:{resultado_extraccion['mime']};base64,{resultado_extraccion['contenido']}"}}
            ]}]

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            max_tokens=1000
        )

        resultado = response.choices[0].message.content.strip()
        resultado = resultado.replace("```json", "").replace("```", "").strip()
        data = json.loads(resultado)
        return jsonify(data)

    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)